from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public" / "scenarios" / "northstar"
OUT.mkdir(parents=True, exist_ok=True)

NAVY = "0B2545"
BLUE = "1457C9"
LIGHT_BLUE = "E8F0FB"
PALE = "F4F7FB"
INK = "182033"
MUTED = "66758E"
GOLD = "C68619"
GREEN = "158365"
RED = "B42318"
RULE = "D7E0EC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_bottom_rule(paragraph, color: str = RULE, size: str = "10") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    text_run.append(text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.extend([fld_char, instr_text, fld_sep, text_run, fld_end])


def apply_doc_style(doc: Document, preset: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25 if preset == "compact_reference_guide" else 1.10

    if preset == "compact_reference_guide":
        settings = {
            "Heading 1": (16, BLUE, 18, 10),
            "Heading 2": (13, BLUE, 14, 7),
            "Heading 3": (12, NAVY, 10, 5),
        }
    else:
        settings = {
            "Heading 1": (16, BLUE, 16, 8),
            "Heading 2": (13, BLUE, 12, 6),
            "Heading 3": (12, NAVY, 8, 4),
        }
    for style_name, (size, color, before, after) in settings.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("NORTHSTAR PRECISION SYSTEMS")
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(NAVY)
        r2 = p.add_run("   |   SYNTHETIC VALIDATION EVIDENCE")
        r2.font.name = "Calibri"
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor.from_string(MUTED)
        set_bottom_rule(p, RULE, "6")

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.add_run("Fictional company • QMSPilot Pilot 1.1 validation • ")
        add_page_field(fp)


def add_memo_masthead(doc: Document, kicker: str, title: str, subtitle: str, metadata: list[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(kicker.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GOLD)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(23)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rl = p.add_run(f"{label}: ")
        rl.bold = True
        rl.font.color.rgb = RGBColor.from_string(INK)
        rv = p.add_run(value)
        rv.font.color.rgb = RGBColor.from_string(INK)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(8)
    set_bottom_rule(rule, BLUE, "12")


def add_label_table(doc: Document, rows: list[tuple[str, str]], compact: bool = False) -> None:
    widths = [1700, 7660] if compact else [2700, 6660]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    for row_idx, (label, value) in enumerate(rows):
        c0, c1 = table.rows[row_idx].cells
        for cell in (c0, c1):
            set_cell_margins(cell, top=60, bottom=60)
        set_cell_shading(c0, LIGHT_BLUE)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor.from_string(NAVY)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor.from_string(INK)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if level:
        p.paragraph_format.left_indent = Inches(0.625)
    p.add_run(text)


def build_document_control() -> Path:
    doc = Document()
    apply_doc_style(doc, "compact_reference_guide")
    add_memo_masthead(
        doc,
        "Controlled procedure",
        "NPS-PR-004 — Document and Data Control",
        "Defines how Northstar approves, distributes, revises, and retires QMS documents and production specifications.",
        [
            ("Revision", "C"),
            ("Effective date", "1 May 2026"),
            ("Process owner", "Quality Systems Manager"),
            ("Approval", "Director of Quality and VP Operations"),
            ("Classification", "Controlled in QMS Vault; printed copies are uncontrolled unless stamped"),
        ],
    )
    doc.add_heading("1. Purpose and scope", level=1)
    doc.add_paragraph(
        "This procedure controls policies, procedures, work instructions, forms, external specifications, engineering drawings, and retained quality records used by Northstar Precision Systems. It applies to the QMS and to both production shifts at the Cedar Falls facility."
    )
    doc.add_heading("2. Control principles", level=1)
    for item in [
        "QMS Vault is the sole approved source for released QMS documents and production drawings.",
        "The legacy shared drive is an archive only. It shall be read-only and shall not be used to release or retrieve production documents after 1 May 2026.",
        "Printed copies are uncontrolled unless they bear a current controlled-copy stamp, revision, issue date, and copy owner.",
        "Document owners must ensure that obsolete electronic copies are removed from points of use within one business day of a new release.",
        "Line supervisors verify the current drawing and work-instruction revision at the start of each shift and record the check on the daily start-up sheet.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3. Roles and authorities", level=1)
    add_label_table(
        doc,
        [
            ("Document owner", "Drafts content, identifies affected users, and confirms technical accuracy."),
            ("Quality Systems", "Assigns document numbers, routes approvals, publishes releases in QMS Vault, and archives superseded versions."),
            ("Approver", "Confirms suitability, resources, and alignment with applicable requirements before release."),
            ("Line supervisor", "Verifies the released revision at point of use and removes obsolete or uncontrolled copies."),
            ("All users", "Use only released documents and report discrepancies immediately."),
        ],
        compact=True,
    )

    doc.add_heading("4. Release and change workflow", level=1)
    for idx, text in enumerate([
        "The owner submits a change request describing the reason, affected processes, and implementation risk.",
        "Quality Systems assigns or confirms the document number and verifies required reviewers.",
        "Technical and process approvers approve the proposed revision before publication.",
        "Quality Systems publishes the approved revision in QMS Vault and records the effective date.",
        "Affected personnel complete required training before independent use of the revised method.",
        "Supervisors remove obsolete copies and record the point-of-use verification within one business day.",
    ], start=1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(text)

    doc.add_heading("5. Records and retention", level=1)
    add_label_table(
        doc,
        [
            ("Document change request", "Seven years after supersession"),
            ("Approval history", "Life of document plus seven years"),
            ("Training acknowledgement", "Employment plus three years"),
            ("Point-of-use revision check", "Two years"),
            ("Obsolete controlled copy log", "Three years"),
        ],
        compact=True,
    )
    doc.add_heading("6. Monitoring", level=1)
    doc.add_paragraph(
        "Quality Systems reviews a quarterly sample of production documents for current revision, approval history, training completion, and removal of obsolete copies. Any use of an obsolete production drawing is entered into the corrective-action system and evaluated for product impact."
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    set_cell_text = p.add_run("END OF CONTROLLED PROCEDURE")
    set_cell_text.bold = True
    set_cell_text.font.color.rgb = RGBColor.from_string(MUTED)
    path = OUT / "02_NPS-PR-004_Document_Control.docx"
    doc.save(path)
    return path


def build_management_review() -> Path:
    doc = Document()
    apply_doc_style(doc, "standard_business_brief")
    add_memo_masthead(
        doc,
        "Management review record",
        "Quality Management Review — June 2026",
        "Leadership review of performance, audit readiness, resources, risks, and improvement actions.",
        [
            ("Meeting date", "30 June 2026"),
            ("Chair", "Elena Ramirez, President"),
            ("Recorder", "Michael Chen, Quality Systems Manager"),
            ("Attendees", "President; VP Operations; Director of Quality; Maintenance Manager; Supply Chain Manager; HR Manager"),
            ("Next review", "28 July 2026"),
        ],
    )
    doc.add_heading("1. Executive status", level=1)
    doc.add_paragraph(
        "Leadership confirmed that the ISO 9001 surveillance audit remains scheduled for 15 September 2026. On-time delivery improved to 94.2%, internal scrap was 2.8%, and customer complaints remained within the annual target. The team agreed that the recent internal-audit findings require focused follow-through."
    )
    doc.add_heading("2. Inputs reviewed", level=1)
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1800, 2100, 2460, 3000])
    table.style = "Table Grid"
    headers = ["Input", "Reported status", "Discussion", "Conclusion"]
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(value)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Internal audit IA-2026-02", "2 nonconformities; 1 observation", "Drawing control and maintenance records require attention.", "Quality and Operations will close items before the surveillance audit."),
        ("Corrective actions", "6 open; 2 reported overdue", "Register migration is still being stabilized.", "Continue biweekly review until backlog is cleared."),
        ("Training", "98% complete", "Revision C rollout is believed substantially complete.", "HR and Operations to finish remaining acknowledgements."),
        ("Supplier performance", "Apex Castings repeat burr issue", "Containment has reduced escapes; root cause is pending.", "Continue supplier monitoring; escalation decision deferred."),
        ("Customer feedback", "One sealing complaint under investigation", "Potential similarity to supplier burr issue was noted.", "Quality will determine whether records should be linked."),
        ("Resources", "Maintenance planner vacancy open", "Backlog may increase during CMMS migration.", "Temporary support may be considered if backlog persists."),
    ]
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_margins(cells[idx])
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(value)

    doc.add_heading("3. Decisions and actions", level=1)
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1400, 4360, 1800, 1800])
    table.style = "Table Grid"
    for idx, value in enumerate(["ID", "Decision / action", "Owner", "Due"]):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(value)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
    set_repeat_table_header(table.rows[0])
    actions = [
        ("MR-26-18", "Resolve the internal-audit items and confirm the evidence package is complete.", "Quality / Operations", "Before audit"),
        ("MR-26-19", "Continue production on HF-1500 and HF-2600 while maintenance records are reconstructed; complete a risk review if concerns emerge.", "Operations", "As needed"),
        ("MR-26-20", "Review the open CAPA backlog at the biweekly leadership meeting.", "Quality team", "Biweekly"),
        ("MR-26-21", "Determine whether Apex Castings requires formal escalation after the next delivery.", "Supply Chain", "Next delivery"),
        ("MR-26-22", "Confirm all affected operators have completed Revision C training.", "HR / Production", "Before audit"),
    ]
    for row_values in actions:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_margins(cells[idx])
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(value)

    doc.add_heading("4. Review conclusion", level=1)
    doc.add_paragraph(
        "Leadership concluded that the QMS remains suitable for the business and that no immediate change to the quality policy is required. The internal-audit response, supplier issue, and remaining training acknowledgements will stay on the monthly watchlist."
    )
    doc.add_heading("5. Approval", level=1)
    add_label_table(
        doc,
        [
            ("Elena Ramirez, President", "Approved electronically 1 July 2026"),
            ("Daniel Brooks, Director of Quality", "Approved electronically 1 July 2026"),
        ],
        compact=True,
    )
    path = OUT / "05_Management_Review_Minutes_2026-06.docx"
    doc.save(path)
    return path


class NorthstarPDF(BaseDocTemplate):
    def __init__(self, filename: str, title: str):
        super().__init__(
            filename,
            pagesize=letter,
            leftMargin=0.68 * inch,
            rightMargin=0.68 * inch,
            topMargin=0.72 * inch,
            bottomMargin=0.62 * inch,
            title=title,
            author="Northstar Precision Systems (synthetic)",
            subject="QMSPilot Pilot 1.1 synthetic validation evidence",
        )
        self.doc_title = title
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="body")
        self.addPageTemplates(PageTemplate(id="northstar", frames=[frame], onPage=self._header_footer))

    def _header_footer(self, canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica-Bold", 8.5)
        canvas.setFillColor(colors.HexColor(f"#{NAVY}"))
        canvas.drawString(self.leftMargin, letter[1] - 0.38 * inch, "NORTHSTAR PRECISION SYSTEMS")
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
        canvas.drawRightString(letter[0] - self.rightMargin, letter[1] - 0.38 * inch, "SYNTHETIC VALIDATION EVIDENCE")
        canvas.setStrokeColor(colors.HexColor(f"#{RULE}"))
        canvas.line(self.leftMargin, letter[1] - 0.46 * inch, letter[0] - self.rightMargin, letter[1] - 0.46 * inch)
        canvas.setFont("Helvetica", 7.5)
        canvas.drawString(self.leftMargin, 0.32 * inch, "Fictional company • No real customer or employee data")
        canvas.drawRightString(letter[0] - self.rightMargin, 0.32 * inch, f"Page {doc.page}")
        canvas.restoreState()


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="KickerNS", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=8.5, leading=10, textColor=colors.HexColor(f"#{GOLD}"), spaceAfter=4))
    styles.add(ParagraphStyle(name="TitleNS", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=22, leading=24, textColor=colors.HexColor(f"#{NAVY}"), alignment=TA_LEFT, spaceAfter=7))
    styles.add(ParagraphStyle(name="SubtitleNS", parent=styles["Normal"], fontName="Helvetica", fontSize=10.5, leading=14, textColor=colors.HexColor(f"#{MUTED}"), spaceAfter=14))
    styles.add(ParagraphStyle(name="H1NS", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor(f"#{BLUE}"), spaceBefore=12, spaceAfter=6, keepWithNext=True))
    styles.add(ParagraphStyle(name="H2NS", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=10.5, leading=13, textColor=colors.HexColor(f"#{NAVY}"), spaceBefore=8, spaceAfter=4, keepWithNext=True))
    styles.add(ParagraphStyle(name="BodyNS", parent=styles["BodyText"], fontName="Helvetica", fontSize=9, leading=12.2, textColor=colors.HexColor(f"#{INK}"), spaceAfter=6))
    styles.add(ParagraphStyle(name="SmallNS", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.8, leading=10.2, textColor=colors.HexColor(f"#{MUTED}"), spaceAfter=3))
    styles.add(ParagraphStyle(name="CalloutNS", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=colors.HexColor(f"#{NAVY}"), backColor=colors.HexColor(f"#{LIGHT_BLUE}"), borderColor=colors.HexColor(f"#{RULE}"), borderWidth=0.5, borderPadding=8, spaceBefore=4, spaceAfter=10))
    styles.add(ParagraphStyle(name="CellNS", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.4, leading=9.4, textColor=colors.HexColor(f"#{INK}")))
    styles.add(ParagraphStyle(name="CellHeadNS", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=7.4, leading=9, textColor=colors.white))
    return styles


S = pdf_styles()


def pdf_title(story, kicker: str, title: str, subtitle: str, metadata: Iterable[tuple[str, str]] = ()):
    story.append(Paragraph(kicker.upper(), S["KickerNS"]))
    story.append(Paragraph(title, S["TitleNS"]))
    story.append(Paragraph(subtitle, S["SubtitleNS"]))
    rows = [[Paragraph(f"<b>{label}</b>", S["CellNS"]), Paragraph(value, S["CellNS"])] for label, value in metadata]
    if rows:
        t = Table(rows, colWidths=[1.45 * inch, 5.55 * inch], hAlign="LEFT")
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{LIGHT_BLUE}")),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{RULE}")),
            ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{RULE}")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 7),
            ("RIGHTPADDING", (0, 0), (-1, -1), 7),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.extend([t, Spacer(1, 10)])


def styled_table(data, widths, header=True, font_size=7.4):
    converted = []
    for r_idx, row in enumerate(data):
        style = S["CellHeadNS"] if header and r_idx == 0 else S["CellNS"]
        converted.append([Paragraph(str(value), style) for value in row])
    t = Table(converted, colWidths=widths, repeatRows=1 if header else 0, hAlign="LEFT")
    commands = [
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("BOX", (0, 0), (-1, -1), 0.55, colors.HexColor(f"#{RULE}")),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{RULE}")),
    ]
    if header:
        commands.append(("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")))
    for r in range(1 if header else 0, len(data)):
        if r % 2 == 0:
            commands.append(("BACKGROUND", (0, r), (-1, r), colors.HexColor(f"#{PALE}")))
    t.setStyle(TableStyle(commands))
    return t


def bullet(text: str):
    return Paragraph(f"• {text}", S["BodyNS"])


def build_scenario_brief() -> Path:
    path = OUT / "00_Northstar_Scenario_Brief.pdf"
    doc = NorthstarPDF(str(path), "Northstar Scenario Brief")
    story = []
    pdf_title(
        story,
        "Pilot 1.1 validation lab",
        "Northstar audit-readiness review",
        "A realistic, completely fictional evidence pack for testing cross-document reasoning, traceability, prioritization, and false-positive control.",
        [
            ("Company", "Northstar Precision Systems (fictional)"),
            ("Scenario date", "16 July 2026"),
            ("Audit target", "ISO 9001 surveillance audit — 15 September 2026"),
            ("Mission", "Prepare a leadership-ready, traceable action brief from the full evidence set"),
        ],
    )
    story.append(Paragraph("How to run the scenario", S["H1NS"]))
    for text in [
        "Open QMSPilot Pilot and load the Northstar mission.",
        "Attach all evidence files in this pack in one review. The files intentionally include mixed formats.",
        "Ask Pilot to cross-reference the records, distinguish direct evidence from inference, and identify both concerns and clean controls.",
        "Do not assume every unusual detail is a finding. Some records are deliberately complete and current.",
        "Compare Pilot's result with the private engineering answer key after the review is complete.",
    ]:
        story.append(bullet(text))
    story.append(Paragraph("Recommended assignment", S["H1NS"]))
    story.append(Paragraph(
        "Review Northstar Precision Systems' audit-readiness evidence. Cross-reference the documents, identify what genuinely requires leadership attention, cite the exact file and record ID for each conclusion, and produce a prioritized action plan with accountable owners, dates, and objective closure evidence. Call out uncertainty where the evidence is incomplete. Do not invent nonconformities.",
        S["CalloutNS"],
    ))
    story.append(Paragraph("Evidence included", S["H1NS"]))
    data = [["File", "Evidence type", "Purpose"]] + [
        ["01_Company_and_QMS_Profile.pdf", "Context", "Business, process, system, and audit context"],
        ["02_NPS-PR-004_Document_Control.docx", "Controlled procedure", "Expected document-control practices"],
        ["03_IA-2026-02_Internal_Audit_Report.pdf", "Audit evidence", "Independent observations and sampled records"],
        ["04_CAPA_Action_Register.xlsx", "Live register", "Corrective-action status, ownership, and verification"],
        ["05_Management_Review_Minutes_2026-06.docx", "Leadership record", "Decisions, commitments, and QMS conclusions"],
        ["06_Training_Matrix.xlsx", "Competence evidence", "Training assignments and completion"],
        ["07_PM_and_Calibration_Records.xlsx", "Asset evidence", "Maintenance, calibration, and use records"],
        ["08_SCAR-2026-007_Supplier_Nonconformance.pdf", "Supplier evidence", "Containment, recurrence, and supplier response"],
        ["09_CC-2026-014_Customer_Complaint.pdf", "Customer evidence", "Complaint investigation and disposition"],
        ["10_Clean_Evidence_Packet.pdf", "Negative controls", "Complete records that should not become findings"],
    ]
    story.append(styled_table(data, [2.3 * inch, 1.25 * inch, 3.45 * inch]))
    story.append(Spacer(1, 10))
    doc.build(story)
    return path


def build_company_profile() -> Path:
    path = OUT / "01_Company_and_QMS_Profile.pdf"
    doc = NorthstarPDF(str(path), "Company and QMS Profile")
    story = []
    pdf_title(
        story,
        "Company profile",
        "Northstar Precision Systems",
        "Fictional manufacturer of precision-machined cooling and sealing assemblies for industrial equipment makers.",
        [
            ("Location", "Cedar Falls, Iowa"),
            ("Employees", "135 employees; two production shifts"),
            ("QMS scope", "Manufacture and final inspection of machined cooling and sealing assemblies"),
            ("Audit date", "ISO 9001 surveillance audit scheduled for 15 September 2026"),
        ],
    )
    story.append(Paragraph("Operating model", S["H1NS"]))
    story.append(Paragraph(
        "Northstar receives aluminum castings from approved suppliers, machines critical sealing and cooling features, performs in-process and final inspection, and ships finished assemblies to industrial original-equipment manufacturers. Product HX-410 accounts for approximately 28% of annual revenue and is produced primarily at machining center M-14 using HF-series hobbing and finishing equipment.",
        S["BodyNS"],
    ))
    story.append(Paragraph("Process and system map", S["H1NS"]))
    data = [
        ["Process", "Owner", "Primary records / systems"],
        ["Contract and customer review", "Commercial Director", "ERP order record; customer specification register"],
        ["Supplier quality", "Supply Chain Manager", "Approved supplier list; incoming inspection; SCAR log"],
        ["Production", "VP Operations", "Released drawing; traveler; start-up verification; equipment record"],
        ["Inspection and release", "Director of Quality", "Inspection plan; gauge record; final-release record"],
        ["Competence", "HR Manager", "Training matrix; qualification record"],
        ["Improvement", "Quality Systems Manager", "Audit log; CAPA register; management review"],
    ]
    story.append(styled_table(data, [1.65 * inch, 1.5 * inch, 3.85 * inch]))
    story.append(Paragraph("Controlled information environment", S["H1NS"]))
    story.append(Paragraph(
        "QMS Vault is Northstar's authoritative repository for released QMS documents and production drawings. Revision C of the document-control procedure made the legacy shared drive archive read-only effective 1 May 2026. QMS Vault lists HX-410 Drawing Revision C as effective 10 June 2026. The shared drive remains available only for historical reference and migration reconciliation.",
        S["CalloutNS"],
    ))
    story.append(Paragraph("Current leadership priorities", S["H1NS"]))
    for text in [
        "Complete the surveillance-audit readiness package by 31 August 2026.",
        "Stabilize preventive-maintenance recordkeeping during the CMMS migration.",
        "Reduce repeat supplier-related burr defects on HX-410 castings.",
        "Move all open corrective actions into one accountable register with objective closure evidence.",
    ]:
        story.append(bullet(text))
    story.append(Paragraph("Performance snapshot", S["H1NS"]))
    data = [
        ["Measure", "June 2026", "Target", "Comment"],
        ["On-time delivery", "94.2%", ">=95%", "Improving from 91.6% in March"],
        ["Internal scrap", "2.8%", "<=2.5%", "HX-410 burr rework is the largest contributor"],
        ["Customer complaints", "3 YTD", "<=6 annual", "One open sealing complaint"],
        ["Training completion", "Reported 98%", ">=98%", "Source is management-review dashboard"],
        ["Open CAPA", "6", "No overdue high-risk items", "Two items reported overdue"],
    ]
    story.append(styled_table(data, [1.55 * inch, 1.15 * inch, 1.55 * inch, 2.75 * inch]))
    story.append(Paragraph("Transition context", S["H1NS"]))
    for text in [
        "QMS Vault and the legacy archive are in a controlled-document transition period.",
        "The maintenance team is reconciling work orders after a CMMS migration.",
        "Leadership dashboards summarize several detailed registers; source records remain the authoritative evidence for individual actions and qualifications.",
        "Northstar expects reviewers to identify contradictions and uncertainty rather than treating every missing field as proof of process failure.",
    ]:
        story.append(bullet(text))
    doc.build(story)
    return path


def build_internal_audit() -> Path:
    path = OUT / "03_IA-2026-02_Internal_Audit_Report.pdf"
    doc = NorthstarPDF(str(path), "Internal Audit Report IA-2026-02")
    story = []
    pdf_title(
        story,
        "Internal audit",
        "IA-2026-02 — Production and Support Processes",
        "Risk-based internal audit covering document control, production, maintenance, competence, supplier quality, and corrective action.",
        [
            ("Audit dates", "23–24 June 2026"),
            ("Lead auditor", "Priya Shah, Internal Auditor (independent of audited work)"),
            ("Areas sampled", "M-14 production, final inspection, maintenance, training, supplier quality, CAPA"),
            ("Result", "2 internal nonconformities; 1 observation; 2 positive controls"),
        ],
    )
    story.append(Paragraph("Audit conclusion", S["H1NS"]))
    story.append(Paragraph(
        "The sampled processes were generally operating, but controls were not consistently demonstrated for released drawings and preventive-maintenance evidence. The audit also found that improvement actions are split across several trackers, reducing leadership visibility. No evidence of shipped nonconforming product was established from the sample; product-impact evaluation remains necessary.",
        S["BodyNS"],
    ))
    story.append(Paragraph("Findings", S["H1NS"]))
    findings = [
        ["ID / class", "Evidence sampled", "Requirement / expected state", "Finding statement"],
        ["IA-26-14\nInternal NC", "At machining center M-14, the start-up packet for lot L26-0619 contained HX-410 Drawing Rev B printed 21 June from the legacy shared drive. QMS Vault showed Rev C effective 10 June. No controlled-copy stamp or completed revision check was present.", "NPS-PR-004 Rev C requires QMS Vault as the sole released source, obsolete copies removed within one business day, and start-of-shift revision verification.", "Northstar did not ensure that the current released drawing was available and verified at point of use."],
        ["IA-26-15\nInternal NC", "PM work orders for HF-1500 (due 31 May) and HF-2600 (due 7 June) were not available in the CMMS or migration folder. Both assets were recorded as operating during June. Staff stated the work may have been completed, but no completion evidence or approved deferral was shown.", "The maintenance program requires scheduled work to be completed, recorded, or formally deferred with risk evaluation.", "Northstar could not demonstrate completion or controlled deferral of required preventive maintenance for two sampled assets."],
        ["IA-26-16\nObservation", "Audit, maintenance, supplier, and training actions were tracked in separate spreadsheets. Meeting notes referenced seven improvement items, but the auditor could not identify one complete register showing owner, date, risk, closure evidence, and effectiveness for all items.", "Management expects material QMS actions to be visible and controlled through the CAPA process.", "Consolidating improvement actions would reduce the risk of overdue or administratively closed work."],
    ]
    story.append(styled_table(findings, [0.75 * inch, 2.35 * inch, 1.85 * inch, 2.05 * inch]))
    story.append(PageBreak())
    story.append(Paragraph("Sample details", S["H1NS"]))
    data = [
        ["Sample", "Record / identifier", "Result"],
        ["Released drawing", "HX-410 Rev C in QMS Vault, effective 10 Jun 2026", "Current controlled revision confirmed"],
        ["Point-of-use packet", "Lot L26-0619, M-14, 21 Jun 2026", "Rev B copy; no revision check"],
        ["Preventive maintenance", "HF-1500 WO-PM-260531; HF-2600 WO-PM-260607", "No completion or deferral evidence"],
        ["Calibration", "Gauge GR-882, certificate CAL-26-188", "Current through 18 Dec 2026; traceability recorded"],
        ["Closed CAPA", "CAPA C-021", "Closure and 60-day effectiveness evidence present"],
        ["Training", "WI-410 Rev C training roll-up", "Dashboard reported 98%; detail matrix not sampled during audit"],
    ]
    story.append(styled_table(data, [1.45 * inch, 3.2 * inch, 2.35 * inch]))
    story.append(Paragraph("Required response", S["H1NS"]))
    for text in [
        "Contain immediate risk and evaluate affected product where applicable.",
        "Record root cause, corrective action, owner, due date, and objective closure evidence in the CAPA register.",
        "Verify effectiveness after implementation; administrative completion alone is not sufficient.",
        "Provide the response to the Internal Audit Lead by 15 July 2026.",
    ]:
        story.append(bullet(text))
    story.append(Paragraph("Auditor note", S["H1NS"]))
    story.append(Paragraph(
        "This report records sampled evidence and does not itself establish product conformity, regulatory compliance, or certification status. Process owners remain responsible for evaluation and disposition.",
        S["SmallNS"],
    ))
    doc.build(story)
    return path


def build_scar() -> Path:
    path = OUT / "08_SCAR-2026-007_Supplier_Nonconformance.pdf"
    doc = NorthstarPDF(str(path), "Supplier Corrective Action Request SCAR-2026-007")
    story = []
    pdf_title(
        story,
        "Supplier quality",
        "SCAR-2026-007 — Burr at coolant-port sealing surface",
        "Supplier corrective-action request for repeat incoming defects on HX-410 castings.",
        [
            ("Supplier", "Apex Castings LLC"),
            ("Opened", "18 June 2026"),
            ("Part / lot", "HX-410 raw casting / AC-061426"),
            ("Quantity", "34 of 200 pieces rejected at incoming inspection"),
            ("Response due", "3 July 2026"),
            ("Status", "Containment accepted; corrective response incomplete"),
        ],
    )
    story.append(Paragraph("Problem description", S["H1NS"]))
    story.append(Paragraph(
        "Incoming inspection found residual burrs on the coolant-port sealing surface. The condition can interfere with gasket seating after machining. The defect is visually similar to SCAR-2026-002, which was closed on 28 March 2026 after the supplier added a final visual check.",
        S["BodyNS"],
    ))
    story.append(Paragraph("Containment", S["H1NS"]))
    data = [
        ["Action", "Owner", "Date", "Evidence / status"],
        ["Quarantine lot AC-061426", "Northstar Incoming Inspection", "18 Jun", "Lot blocked; 34 rejected"],
        ["100% sort of Northstar stock", "Apex Resident Technician", "19 Jun", "612 pieces screened; 11 additional pieces rejected"],
        ["Inspect next three supplier shipments", "Northstar Supplier Quality", "From 20 Jun", "Two shipments accepted after enhanced inspection"],
    ]
    story.append(styled_table(data, [2.0 * inch, 1.6 * inch, 0.8 * inch, 2.6 * inch]))
    story.append(Paragraph("Supplier response received 1 July 2026", S["H1NS"]))
    add = [
        ["Response element", "Supplier statement", "Northstar review"],
        ["Cause", "Operator inconsistency during manual deburr and final visual inspection.", "No process evidence or causal analysis attached; does not explain recurrence after SCAR-2026-002."],
        ["Correction", "Reinstruct operators and remind inspectors to check all coolant ports.", "Training roster not provided."],
        ["Corrective action", "Add a quality alert at the deburr station.", "Implementation date and control-plan change not provided."],
        ["Effectiveness", "No additional defects expected.", "No sampling plan, duration, metric, or acceptance criterion defined."],
    ]
    story.append(styled_table(add, [1.2 * inch, 2.65 * inch, 3.15 * inch]))
    story.append(Paragraph("Northstar disposition", S["H1NS"]))
    story.append(Paragraph(
        "Containment remains active. Supplier Quality requested a revised response with verified cause, objective implementation evidence, and an effectiveness plan. No formal supplier escalation level or sourcing decision has been approved. The SCAR record is not currently linked to customer complaint CC-2026-014 or to a dedicated internal CAPA.",
        S["CalloutNS"],
    ))
    doc.build(story)
    return path


def build_complaint() -> Path:
    path = OUT / "09_CC-2026-014_Customer_Complaint.pdf"
    doc = NorthstarPDF(str(path), "Customer Complaint CC-2026-014")
    story = []
    pdf_title(
        story,
        "Customer quality",
        "CC-2026-014 — Seal leakage after assembly",
        "Customer complaint investigation record for HX-410 finished assemblies.",
        [
            ("Customer", "Titan Thermal Equipment"),
            ("Received", "22 June 2026"),
            ("Product / lot", "HX-410 / Northstar lot L26-0619"),
            ("Quantity", "6 units reported; 2 returned for evaluation"),
            ("Status", "Open — interim response sent 26 June 2026"),
        ],
    )
    story.append(Paragraph("Customer statement", S["H1NS"]))
    story.append(Paragraph(
        "Titan Thermal reported coolant leakage during end-of-line pressure testing. The customer observed incomplete gasket seating at one coolant port and requested containment, replacement units, and a cause-and-corrective-action response.",
        S["BodyNS"],
    ))
    story.append(Paragraph("Northstar investigation", S["H1NS"]))
    data = [
        ["Evidence", "Result"],
        ["Returned unit examination", "Both returned units showed a raised edge / burr witness at the coolant-port sealing surface."],
        ["Lot trace", "Northstar lot L26-0619 used Apex Castings supplier lot AC-061426."],
        ["Production record", "Machined at M-14 on 21 June 2026; in-process dimensions within recorded limits."],
        ["Final inspection", "Visual and dimensional acceptance recorded. Final inspection used torque wrench TW-118 for fixture setup."],
        ["Drawing review", "Investigator noted that the production packet contained HX-410 Rev B; impact of the Rev B / Rev C difference has not been documented."],
        ["Similar records", "The complaint owner noted a possible similarity to recent incoming burr defects but did not record a SCAR or CAPA reference."],
    ]
    story.append(styled_table(data, [1.55 * inch, 5.45 * inch]))
    story.append(Paragraph("Containment and response", S["H1NS"]))
    for text in [
        "Northstar placed remaining finished goods from lot L26-0619 on hold and initiated 100% visual inspection.",
        "Replacement units were shipped to Titan Thermal on 24 June 2026.",
        "The customer accepted the interim containment and requested the final corrective-action response by 29 July 2026.",
        "No documented decision has been made on whether the complaint, SCAR-2026-007, IA-26-14, or an internal CAPA should share one root-cause investigation.",
    ]:
        story.append(bullet(text))
    story.append(Paragraph("Open questions", S["H1NS"]))
    story.append(Paragraph(
        "The evidence supports a plausible relationship among the customer complaint, the supplier burr condition, and the point-of-use drawing issue, but causation has not been established. A controlled investigation must distinguish contributing factors and document product impact.",
        S["CalloutNS"],
    ))
    doc.build(story)
    return path


def build_clean_packet() -> Path:
    path = OUT / "10_Clean_Evidence_Packet.pdf"
    doc = NorthstarPDF(str(path), "Clean Evidence Packet")
    story = []
    pdf_title(
        story,
        "Negative controls",
        "Complete and current sampled records",
        "These records are included to test whether Pilot can recognize satisfactory evidence without manufacturing findings.",
        [
            ("Packet date", "10 July 2026"),
            ("Prepared by", "Northstar Quality Systems"),
            ("Records", "Gauge calibration; closed CAPA; approved supplier certificate"),
        ],
    )
    story.append(Paragraph("Record 1 — Calibration certificate CAL-26-188", S["H1NS"]))
    data = [
        ["Field", "Value"],
        ["Instrument", "Digital bore gauge GR-882"],
        ["Calibration date", "18 June 2026"],
        ["Next due", "18 December 2026"],
        ["Result", "As found / as left within specification"],
        ["Standards", "Reference set RS-44; traceable to NIST through certificate LAB-889102"],
        ["Laboratory", "Midwest Metrology Services; ISO/IEC 17025 scope verified at time of service"],
        ["Approval", "Electronically approved by L. Novak, Metrology Lead, 19 June 2026"],
    ]
    story.append(styled_table(data, [1.65 * inch, 5.35 * inch]))
    story.append(Paragraph("Record 2 — CAPA C-021 closure and effectiveness", S["H1NS"]))
    data = [
        ["Element", "Evidence"],
        ["Problem", "Incorrect label applied to two internal transfer bins; no customer shipment affected."],
        ["Cause", "Label template selected by description instead of controlled part number."],
        ["Action", "Template library keyed to part number; obsolete templates removed; users trained."],
        ["Closure evidence", "System change CR-2026-041; access log; training records for 12 affected users."],
        ["Effectiveness", "Sixty-day sample of 240 transfer-bin labels found zero recurrence; reviewed 12 June 2026."],
        ["Approval", "Closed effective by Quality Systems Manager and Warehouse Manager on 13 June 2026."],
    ]
    story.append(styled_table(data, [1.65 * inch, 5.35 * inch]))
    story.append(PageBreak())
    story.append(Paragraph("Record 3 — Approved supplier certificate", S["H1NS"]))
    data = [
        ["Field", "Value"],
        ["Supplier", "Blue Ridge Anodizing"],
        ["Approval scope", "Type II and Type III anodize for non-HX-410 product families"],
        ["Certificate", "QMS certificate BR-9001-2451, valid through 31 March 2027"],
        ["Performance review", "98.7% acceptance; zero late deliveries; no open SCAR"],
        ["Northstar approval", "Annual review completed 4 April 2026; next review due April 2027"],
    ]
    story.append(styled_table(data, [1.65 * inch, 5.35 * inch]))
    story.append(Paragraph("Validation note", S["H1NS"]))
    story.append(Paragraph(
        "A strong analysis may reference these records as evidence of effective controls or as contrast with weaker records elsewhere. It should not turn them into corrective-action findings without contradictory evidence.",
        S["CalloutNS"],
    ))
    doc.build(story)
    return path


def main() -> None:
    files = [
        build_scenario_brief(),
        build_company_profile(),
        build_document_control(),
        build_internal_audit(),
        build_management_review(),
        build_scar(),
        build_complaint(),
        build_clean_packet(),
    ]
    print("Generated:")
    for path in files:
        print(path.relative_to(ROOT))


if __name__ == "__main__":
    main()
